import streamlit as st
from openai import OpenAI
import io
from docx import Document

st.set_page_config(page_title="ZenMux AI 助手", page_icon="🦀", layout="wide")

# 初始化云端会话状态（阅后即焚，不存本地）
if "api_key" not in st.session_state:
    st.session_state.api_key = ""
if "messages" not in st.session_state:
    st.session_state.messages = []

def generate_word_doc(messages):
    doc = Document()
    doc.add_heading('ZenMux AI 对话记录', 0)
    for msg in messages:
        if msg["role"] == "system": continue
        role_name = "🧑‍💻 我" if msg["role"] == "user" else "🤖 AI"
        doc.add_heading(role_name, level=2)
        doc.add_paragraph(msg["content"])
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

with st.sidebar:
    st.header("⚙️ 核心配置")
    st.info("💡 提示：为保护隐私，云端版本不会记录您的 API Key 与对话内容，刷新网页后自动清空。")
    
    api_key = st.text_input(
        "🔑 ZenMux API Key", 
        type="password", 
        value=st.session_state.api_key,
        placeholder="输入 API Key...",
    )
    st.session_state.api_key = api_key

    model_name = st.selectbox(
        "🧠 选择模型", 
        ["anthropic/claude-opus-4.6", "anthropic/claude-sonnet-4.6", "anthropic/claude-haiku-4.5"]
    )
    
    st.divider()
    
    st.header("📄 附加资料区")
    uploaded_file = st.file_uploader("上传参考文档 (txt, md, csv)", type=['txt', 'md', 'csv'])
    file_content = ""
    if uploaded_file is not None:
        try:
            file_content = uploaded_file.getvalue().decode("utf-8")
            st.success(f"✅ [{uploaded_file.name}] 已载入")
        except Exception as e:
            st.error("❌ 解析失败")

    st.divider()
    
    st.header("📦 导出与清理")
    if st.button("🗑️ 清空当前对话", use_container_width=True):
        st.session_state.messages = []
        st.rerun()
        
    if st.session_state.messages:
        c1, c2 = st.columns(2)
        with c1:
            txt_content = "".join([f"{'我' if m['role']=='user' else 'AI'}:\n{m['content']}\n\n{'-'*40}\n\n" for m in st.session_state.messages if m["role"]!="system"])
            st.download_button("📥 TXT", txt_content, "AI对话.txt", "text/plain")
        with c2:
            st.download_button("📥 Word", generate_word_doc(st.session_state.messages), "AI对话.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.title("🤖 ZenMux 多模态智能助手")

if not st.session_state.api_key:
    st.warning("👈 请先在左侧边栏输入您的 API Key。")
    st.stop()

client = OpenAI(base_url="https://zenmux.ai/api/v1", api_key=st.session_state.api_key)

for message in st.session_state.messages:
    if message["role"] == "system": continue
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

if prompt := st.chat_input(f"想问点什么？(当前使用: {model_name})"):
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)
        
    api_messages = []
    if file_content:
        api_messages.append({"role": "system", "content": f"基于以下文件内容回答：\n\n{file_content}"})
    api_messages.extend([{"role": m["role"], "content": m["content"]} for m in st.session_state.messages])

    with st.chat_message("assistant"):
        try:
            stream = client.chat.completions.create(model=model_name, messages=api_messages, stream=True)
            response = st.write_stream(stream)
            st.session_state.messages.append({"role": "assistant", "content": response})
            st.rerun()
        except Exception as e:
            st.error(f"请求失败: {e}")
